from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "technical_audit_final_review.md"
OUTPUT = ROOT / "docs" / "CSMS_Technical_Audit_Final_Review.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF6FC"
PALE_YELLOW = "FFF2CC"
PALE_GREEN = "E2F0D9"
LIGHT_GRAY = "F2F2F2"
DARK_GRAY = "404040"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("<br>", "\n")
    text = text.replace("<br/>", "\n")
    return text.strip()


def add_runs_with_inline_code(paragraph, text: str, size: int | None = None) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(31, 78, 121)
            if size:
                run.font.size = Pt(size)
        else:
            run = paragraph.add_run(part)
            if size:
                run.font.size = Pt(size)


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, 15000)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.size = Pt(10.5)
    if body:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_runs_with_inline_code(p2, body, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("CSMS Technical Audit")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mini Project Final Review Handbook")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(64, 64, 64)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Java + JavaFX + OracleDB + JDBC + SMTP + iText")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(64, 64, 64)

    add_callout(
        doc,
        "How to use this document tomorrow",
        "First revise the Quick Answers. Then read each workflow section. For viva/code review, explain the flow in this order: FXML view -> Java controller -> DAO/service -> OracleDB/file/email/PDF output.",
        PALE_GREEN,
    )

    quick = [
        ("UI framework", "JavaFX FXML. No Swing JTabbedPane. Navigation uses sidebar buttons and FXMLLoader."),
        ("Database", "OracleDB through JDBC using ojdbc8.jar. DAO classes use PreparedStatement and ResultSet."),
        ("Email", "SMTP through Gmail host smtp.gmail.com, port 587, STARTTLS, Jakarta Mail API with Angus Mail implementation."),
        ("PDF storage", "Assignments, submissions, and question papers are files in folders. Metadata is in OracleDB."),
        ("BLOB usage", "Profile images and school profile image are stored as Oracle BLOBs."),
        ("Chat", "Database-backed chat through COMMUNICATION and COMMUNICATION_READ_STATE. No WebSockets."),
        ("Reports", "Performance PDFs are generated on demand using iText libraries."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Review Topic"
    table.rows[0].cells[1].text = "Short Answer"
    style_table(table, [2400, 12600], header_fill=BLUE)
    for left, right in quick:
        row = table.add_row()
        row.cells[0].text = left
        row.cells[1].text = right
    style_table(table, [2400, 12600], header_fill=BLUE)
    doc.add_page_break()


def style_table(table, widths: list[int] | None = None, header_fill: str = BLUE) -> None:
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 120, 120, 120, 120)
            if widths and c_idx < len(widths):
                set_cell_width(cell, widths[c_idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8.5 if len(row.cells) >= 4 else 9.5)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)


def add_table_chunk(doc: Document, rows: list[list[str]], widths: list[int] | None) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    for i, text in enumerate(rows[0]):
        table.rows[0].cells[i].text = text
    for row_values in rows[1:]:
        row = table.add_row()
        for i in range(cols):
            row.cells[i].text = row_values[i] if i < len(row_values) else ""
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def add_markdown_table(doc: Document, block: list[str]) -> None:
    rows: list[list[str]] = []
    for line in block:
        if re.match(r"^\|\s*-", line):
            continue
        cells = [clean_inline(c) for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    cols = len(rows[0])
    if cols == 2:
        widths = [3800, 11200]
        chunk_size = 14
    elif cols == 3:
        widths = [2900, 3500, 8600]
        chunk_size = 10
    elif cols == 4:
        widths = [2800, 3600, 3100, 5500]
        chunk_size = 8
    elif cols == 5:
        widths = [2600, 2700, 4400, 3400, 3800]
        chunk_size = 6
    else:
        widths = None
        chunk_size = 8

    header, body = rows[0], rows[1:]
    if len(body) <= chunk_size:
        add_table_chunk(doc, rows, widths)
        return

    for idx in range(0, len(body), chunk_size):
        if idx:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("Table continued")
            r.italic = True
            r.font.color.rgb = RGBColor(96, 96, 96)
            r.font.size = Pt(9)
        add_table_chunk(doc, [header] + body[idx : idx + chunk_size], widths)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 9)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor(31, 78, 121 if level <= 2 else 64)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18 if level == 0 else 0.35)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("- ")
    add_runs_with_inline_code(p, clean_inline(text), 10)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(2)
    add_runs_with_inline_code(p, clean_inline(text), 10)


def build_doc() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    add_cover(doc)

    pending_table: list[str] = []
    in_code_block = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_callout(doc, "Code / command reference", "\n".join(code_lines), PALE_BLUE)
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            pending_table.append(line)
            continue
        if pending_table:
            add_markdown_table(doc, pending_table)
            pending_table = []

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            add_heading(doc, clean_inline(stripped[2:]), 1)
        elif stripped.startswith("## "):
            title = clean_inline(stripped[3:])
            if title in {"1. Component Mapping", "2. Workflow Logic", "3. External JAR Mystery", "4. Major Class Responsibilities", "5. Coding Workflow: How The Project Works End To End"}:
                doc.add_page_break()
            add_heading(doc, title, 1)
        elif stripped.startswith("### "):
            add_heading(doc, clean_inline(stripped[4:]), 2)
        elif stripped.startswith(">"):
            add_callout(doc, "Say this in review", clean_inline(stripped.lstrip("> ")), PALE_YELLOW)
        elif stripped in {"Why this design:", "Important answer:", "What to say if asked about handshake:", "Short review explanation:", "Use this as your explanation flow in review:", "Simple MVC-style explanation:"}:
            add_callout(doc, stripped.rstrip(":"), "", PALE_BLUE)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            add_numbered(doc, stripped)
        else:
            p = doc.add_paragraph()
            add_runs_with_inline_code(p, clean_inline(stripped), 10.2)

    if pending_table:
        add_markdown_table(doc, pending_table)

    # Footer-like closing study note.
    doc.add_page_break()
    add_heading(doc, "Final One-Minute Explanation", 1)
    add_callout(
        doc,
        "Memorize this flow",
        "FXML defines the screen. The Java controller handles button clicks and table setup. The controller calls a service or DAO. DAO uses JDBC and Oracle SQL. Results are converted into model classes and displayed in JavaFX tables/charts. Files and PDFs are handled by service classes, while email is handled by SMTP through Jakarta Mail.",
        PALE_GREEN,
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
