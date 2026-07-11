import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SQL_PATH = ROOT / "database" / "csms_full_setup_with_data.sql"
DOC_PATH = ROOT / "docs" / "csms_table_structures.docx"
MD_PATH = ROOT / "docs" / "csms_table_structures.md"


DESCRIPTIONS = {
    "ROLE_ID": "Unique role ID",
    "ROLE_NAME": "Role name",
    "USER_ID": "Related user ID",
    "NAME": "Full name of the user",
    "EMAIL": "User email address",
    "PASSWORD_HASH": "Encrypted user password",
    "PHONE": "User contact number",
    "CREATED_AT": "Account creation date",
    "IS_ACTIVE": "User active status",
    "CLASS_ID": "Related class ID",
    "CLASS_NAME": "Class name",
    "SECTION": "Class section",
    "ACADEMIC_YEAR": "Academic year",
    "SUBJECT_ID": "Related subject ID",
    "SUBJECT_NAME": "Subject name",
    "TEACHER_ID": "Related teacher ID",
    "QUALIFICATION": "Teacher qualification",
    "EXPERIENCE": "Teaching experience in years",
    "STUDENT_ID": "Related student ID",
    "DOB": "Student date of birth",
    "GENDER": "Student gender",
    "CONDUCT": "Student conduct grade",
    "CONDUCT_REMARKS": "Remarks about student conduct",
    "PARENT_ID": "Parent user ID",
    "RELATION": "Relationship with student",
    "QP_ID": "Related question paper ID",
    "EXAM_TYPE": "Type or title of exam",
    "EXAM_DATE": "Date of examination",
    "MAX_MARKS": "Maximum marks for exam",
    "EXAM_DESCRIPTION": "Description of the exam",
    "CREATED_BY_TEACHER_ID": "Teacher who created the exam",
    "QUESTION_ID": "Unique question bank ID",
    "TITLE": "Title of assignment or question paper",
    "ORIGINAL_FILE_NAME": "Original uploaded file name",
    "UPLOADED_AT": "File uploaded date",
    "ASSIGNMENT_ID": "Related assignment ID",
    "DESCRIPTION": "Assignment description",
    "DUE_DATE": "Assignment due date",
    "MARK_ID": "Unique marks entry ID",
    "MARKS_OBTAINED": "Marks obtained by student",
    "SUBMISSION_ID": "Unique submission ID",
    "SUBMITTED_ON": "Submission date",
    "MARKS": "Marks awarded for submission",
    "ATTENDANCE_ID": "Unique attendance ID",
    "ATTENDANCE_DATE": "Attendance date",
    "SESSION_TYPE": "Attendance session type",
    "STATUS": "Current status",
    "LEAVE_REASON": "Reason for leave",
    "APPROVAL_STATUS": "Leave approval status",
    "APPROVED_BY": "User who approved leave",
    "SESSION_ID": "Unique counselling session ID",
    "COUNSELLOR_ID": "Assigned counsellor user ID",
    "SESSION_DATE": "Counselling session date",
    "NOTES": "Counselling notes",
    "CATEGORY": "Counselling category",
    "MESSAGE_ID": "Unique message ID",
    "SENDER_ID": "Message sender user ID",
    "RECEIVER_ID": "Message receiver user ID",
    "MESSAGE_TEXT": "Message content",
    "SENT_AT": "Message sent timestamp",
    "VIEWER_ID": "User viewing the conversation",
    "PARTNER_ID": "Conversation partner user ID",
    "LAST_SEEN_MESSAGE_ID": "Last message seen by viewer",
    "LAST_SEEN_AT": "Last seen timestamp",
    "LOG_ID": "Unique login audit ID",
    "LOGIN_TIME": "Login time",
    "LOGOUT_TIME": "Logout time",
    "IP_ADDRESS": "Login IP address",
    "SETTING_KEY": "Unique setting key",
    "SETTING_VALUE": "Stored setting value",
    "IMAGE_DATA": "Binary image data",
    "MIME_TYPE": "Image MIME type",
    "UPDATED_AT": "Last updated date",
    "IMAGE_ID": "Unique school image ID",
}


def nice_name(name):
    return name.replace("_", " ").title()


def split_sql_items(body):
    items = []
    current = []
    depth = 0
    in_quote = False
    for char in body:
        if char == "'":
            in_quote = not in_quote
        if not in_quote:
            if char == "(":
                depth += 1
            elif char == ")":
                depth -= 1
            elif char == "," and depth == 0:
                item = "".join(current).strip()
                if item:
                    items.append(item)
                current = []
                continue
        current.append(char)
    item = "".join(current).strip()
    if item:
        items.append(item)
    return items


def parse_type(raw_type):
    match = re.match(r"([A-Z0-9]+)(?:\(([^)]+)\))?", raw_type)
    if not match:
        return raw_type, ""
    sql_type = match.group(1)
    width = match.group(2) or ""
    type_map = {
        "NUMBER": "Number",
        "VARCHAR2": "String",
        "DATE": "Date",
        "TIMESTAMP": "Timestamp",
        "BLOB": "BLOB",
    }
    return type_map.get(sql_type, sql_type.title()), width


def parse_tables(sql):
    pattern = re.compile(r"CREATE TABLE\s+([A-Z_]+)\s*\((.*?)\);", re.S)
    tables = []
    for table_name, body in pattern.findall(sql):
        items = split_sql_items(body)
        columns = []
        table_constraints = []
        for item in items:
            normalized = " ".join(item.split())
            if normalized.startswith("CONSTRAINT ") or normalized.startswith("PRIMARY KEY"):
                table_constraints.append(normalized)
                continue
            parts = normalized.split(" ", 2)
            if len(parts) < 2:
                continue
            col_name = parts[0]
            raw_type = parts[1]
            rest = parts[2] if len(parts) > 2 else ""
            col_type, width = parse_type(raw_type)
            constraints = []
            for token in ["PRIMARY KEY", "UNIQUE", "NOT NULL"]:
                if token in rest:
                    constraints.append(token)
            default_match = re.search(r"DEFAULT\s+(.+?)(?:\s+NOT NULL|\s+UNIQUE|\s+PRIMARY KEY|$)", rest)
            if default_match:
                constraints.append(f"DEFAULT {default_match.group(1).strip()}")
            check_match = re.search(r"CHECK\s*\((.+)\)", rest)
            if check_match:
                constraints.append(f"CHECK ({check_match.group(1).strip()})")
            ref_match = re.search(r"REFERENCES\s+([A-Z_]+\([A-Z_]+\))", rest)
            if ref_match:
                constraints.append(f"FOREIGN KEY -> {ref_match.group(1)}")
            columns.append(
                {
                    "name": col_name,
                    "type": col_type,
                    "width": width,
                    "constraints": constraints,
                    "description": DESCRIPTIONS.get(col_name, f"{nice_name(col_name)} value"),
                }
            )

        for constraint in table_constraints:
            pk = re.search(r"PRIMARY KEY\s*\(([^)]+)\)", constraint)
            if pk:
                for col in [c.strip() for c in pk.group(1).split(",")]:
                    add_constraint(columns, col, "PRIMARY KEY")
            uq = re.search(r"UNIQUE\s*\(([^)]+)\)", constraint)
            if uq:
                label = f"UNIQUE: {uq.group(1)}"
                for col in [c.strip() for c in uq.group(1).split(",")]:
                    add_constraint(columns, col, label)
            fk = re.search(r"FOREIGN KEY\s*\(([^)]+)\)\s+REFERENCES\s+([A-Z_]+\([A-Z_]+\))", constraint)
            if fk:
                label = f"FOREIGN KEY -> {fk.group(2)}"
                for col in [c.strip() for c in fk.group(1).split(",")]:
                    add_constraint(columns, col, label)
            chk = re.search(r"CHECK\s*\((.+)\)", constraint)
            if chk:
                check_expr = chk.group(1)
                for column in columns:
                    if re.search(rf"\b{re.escape(column['name'])}\b", check_expr):
                        add_constraint(columns, column["name"], f"CHECK ({check_expr})")

        tables.append({"name": table_name, "columns": columns})
    return tables


def add_constraint(columns, col_name, label):
    for column in columns:
        if column["name"] == col_name and label not in column["constraints"]:
            column["constraints"].append(label)


def set_cell_border(cell, size="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, value, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_column_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def column_description(column):
    constraints = " ".join(column["constraints"])
    if "PRIMARY KEY" in constraints and "FOREIGN KEY" not in constraints:
        return f"Unique {nice_name(column['name']).lower()}"
    if "FOREIGN KEY" in constraints:
        return f"Related {nice_name(column['name']).lower()}"
    return column["description"]


def add_paragraph(doc, text_value, size=14, bold=False, before=0, after=8):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text_value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    return paragraph


def add_table(doc, table, index):
    add_paragraph(doc, f"3.6.{index} {nice_name(table['name'])} Structure:", size=15, bold=True, before=4, after=16)
    add_paragraph(doc, f"Table 3.{index}: {nice_name(table['name'])} Description", size=15, bold=True, before=2, after=18)

    word_table = doc.add_table(rows=1, cols=5)
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = False
    widths = [3.2, 2.5, 2.1, 5.5, 5.0]
    headers = ["Attribute\nName", "Type", "Width", "Constraint(s)", "Description"]
    for cell, header, width in zip(word_table.rows[0].cells, headers, widths):
        set_column_width(cell, width)
        set_cell_text(cell, header, bold=True, size=14)
        shade_cell(cell, "F2F2F2")
        set_cell_border(cell)

    for column in table["columns"]:
        row = word_table.add_row()
        values = [
            column["name"].lower(),
            column["type"],
            column["width"],
            ",\n".join(column["constraints"]) if column["constraints"] else "-",
            column_description(column),
        ]
        for cell, value, width in zip(row.cells, values, widths):
            set_column_width(cell, width)
            align = WD_ALIGN_PARAGRAPH.LEFT if cell is row.cells[4] else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, value, size=12, align=align)
            set_cell_border(cell)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def write_markdown(tables):
    lines = ["# CSMS Table Structures", ""]
    for index, table in enumerate(tables, start=1):
        lines.extend(
            [
                f"## 3.6.{index} {nice_name(table['name'])} Structure:",
                "",
                f"**Table 3.{index}: {nice_name(table['name'])} Description**",
                "",
                "| Attribute Name | Type | Width | Constraint(s) | Description |",
                "|---|---|---:|---|---|",
            ]
        )
        for column in table["columns"]:
            constraints = "<br>".join(column["constraints"]) if column["constraints"] else "-"
            lines.append(
                f"| {column['name'].lower()} | {column['type']} | {column['width']} | {constraints} | {column_description(column)} |"
            )
        lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def main():
    sql = SQL_PATH.read_text(encoding="utf-8")
    tables = parse_tables(sql)
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_paragraph(doc, "3.6 CSMS Database Table Structures", size=18, bold=True, after=14)
    add_paragraph(
        doc,
        "The following tables describe all database entities used in the Centralized School Management System.",
        size=12,
        after=18,
    )
    for index, table in enumerate(tables, start=1):
        add_table(doc, table, index)

    doc.save(DOC_PATH)
    write_markdown(tables)
    print(f"Generated {DOC_PATH}")
    print(f"Generated {MD_PATH}")
    print(f"Tables: {len(tables)}")


if __name__ == "__main__":
    main()
